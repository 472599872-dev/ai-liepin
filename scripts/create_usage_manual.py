from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "猎聘招聘智能体使用手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 45)
MUTED = RGBColor(91, 102, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, *, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_para_tokens(paragraph, *, before: float = 0, after: float = 6, line_spacing: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def paragraph_border_left(paragraph, color: str = "2E74B5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def paragraph_shading(paragraph, fill: str = CALLOUT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_tokens(p, before=92, after=10, line_spacing=1.15)
    r = p.add_run("猎聘招聘智能体使用手册")
    set_run_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_tokens(p, after=18, line_spacing=1.15)
    r = p.add_run("本地工作台 / Windows 便携版 / 猎头日常操作参考")
    set_run_font(r, size=13, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [dxa(1.55), dxa(4.95)])
    set_borders(meta)
    rows = [
        ("版本", "v1.0，2026-06-12"),
        ("适用对象", "使用猎聘找人、批量分析简历、维护候选人清单的猎头用户"),
        ("交付形态", "本地桌面应用；Windows 打包后可双击运行"),
        ("重要边界", "遵守平台规则，控制频率；遇到登录、验证码、权益、页面异常时由系统提醒人工处理"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), size=10.5, bold=True, color=DARK_BLUE)
        set_run_font(row.cells[1].paragraphs[0].add_run(value), size=10.5, color=INK)

    p = doc.add_paragraph()
    set_para_tokens(p, before=18, after=8)
    paragraph_shading(p, CALLOUT)
    paragraph_border_left(p)
    r = p.add_run("一句话流程：维护账号和岗位 -> 解析 JD 并确认猎聘条件 -> 创建任务 -> 系统按队列搜索、打开候选人、抓简历、评分、达阈值后沟通 -> 在候选人清单复盘。")
    set_run_font(r, size=11, bold=True, color=INK)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, *, bold_prefix: str = "") -> None:
    p = doc.add_paragraph()
    set_para_tokens(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    set_para_tokens(p, after=4)
    for run in p.runs:
        set_run_font(run, size=11, color=INK)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    set_para_tokens(p, after=4)
    for run in p.runs:
        set_run_font(run, size=11, color=INK)


def add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    set_para_tokens(p, before=4, after=8)
    paragraph_shading(p, CALLOUT)
    paragraph_border_left(p)
    r = p.add_run(f"{label}：")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, [dxa(x) for x in widths_in])
    set_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.2, color=INK)
    set_table_geometry(table, [dxa(x) for x in widths_in])
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11, color=INK)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, size=11, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "猎聘招聘智能体使用手册"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_tokens(header, after=0, line_spacing=1.0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_tokens(footer, after=0, line_spacing=1.0)
    r = footer.add_run("本地资料，请勿外传包含 .env、data、profiles 的完整包")
    set_run_font(r, size=8.5, color=MUTED)

    doc.core_properties.title = "猎聘招聘智能体使用手册"
    doc.core_properties.subject = "本地招聘自动化工作台操作手册"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "猎聘, 招聘智能体, 使用手册, Windows"


def build() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. 快速开始", 1)
    add_body(doc, "这份手册面向日常使用，不解释源码实现。使用者只需要知道账号、岗位、任务、候选人、调试日志这五个入口即可。")
    add_number(doc, "打开 Windows 成品目录，双击 LiepinRecruitingAgent.exe。")
    add_number(doc, "进入账号页，选择或新增猎聘账号；右侧猎聘页面会使用该账号独立的本地 Profile。")
    add_number(doc, "进入岗位页，新增岗位并粘贴 JD；在表单页点击解析 JD，确认哪些猎聘查询条件要启用。")
    add_number(doc, "进入任务页，选择账号、岗位、目标人数、打招呼阈值、是否 AI 评分、是否自动沟通。")
    add_number(doc, "运行任务后，系统会按队列顺序处理；需要人工介入时看提醒和调试日志。")
    add_note(doc, "默认建议", "首次跑新岗位时先保持 Dry-run，确认搜索条件、评分结果和话术都符合预期后，再开启真实发送。")

    add_heading(doc, "2. 界面总览", 1)
    add_matrix(
        doc,
        ["区域", "用途", "高频操作"],
        [
            ["左侧账号", "维护账号名称、手机号、密码和登录状态。每个账号有独立 Profile。", "选择账号、编辑密码、检查登录状态"],
            ["左侧岗位", "维护 JD 和猎聘查询条件。解析 JD 是岗位表单页功能。", "新增岗位、编辑岗位、解析 JD、确认条件"],
            ["左侧任务", "编排批量执行计划。任务按排序顺序逐个跑。", "新增任务、运行队列、终止当前任务、调整顺序"],
            ["左侧候选人", "按岗位查看候选人资产。支持排序、评分筛选、状态筛选和淘汰。", "查看详情、打开猎聘简历、看全文快照"],
            ["右侧工作区", "内嵌猎聘网页。默认只保留网页，调试内容收在调试弹窗里。", "观察页面、人工登录、处理验证码、确认异常"],
        ],
        [1.25, 2.5, 2.5],
    )
    add_body(doc, "右侧猎聘页面是实际执行现场。系统点击、填条件、打开候选人、沟通窗口都会在这里发生；调试按钮只在需要看过程、录制操作或单测分页时打开。")

    add_heading(doc, "3. 账号维护与登录", 1)
    add_body(doc, "账号列表直接维护账号和密码。选择某个账号后，系统会自动切到密码登录页并填入对应账号密码；登录态保存在该账号自己的 Profile 目录中。")
    add_bullet(doc, "同一个人维护多个账号时，不需要反复退出登录；系统运行任务前会切换到任务绑定的账号 Profile。")
    add_bullet(doc, "从候选人清单打开猎聘简历时，优先使用当时抓取该简历的账号，而不是当前随手选中的账号。")
    add_bullet(doc, "遇到验证码、短信、安全验证时，系统会停下并提示人工处理；处理完成后再继续任务。")
    add_note(doc, "登录态", "profiles/app/account_<账号ID> 保存登录态，data/app.db 保存账号记录。迁移机器时要一起迁移，但这些目录含敏感信息。")

    add_heading(doc, "4. 岗位与 JD 配置", 1)
    add_body(doc, "岗位页默认是列表。新增或编辑时才打开表单；JD、岗位标题、猎聘查询条件、评分阈值、默认话术都在表单中维护。")
    add_matrix(
        doc,
        ["字段", "说明", "使用建议"],
        [
            ["岗位标题", "内部识别名称，也会影响任务名称和候选人归档。", "写清楚岗位方向，避免多个岗位混淆"],
            ["岗位 JD", "用于解析猎聘条件、评分、生成补充话术。", "尽量包含职责、必备项、加分项、行业场景"],
            ["猎聘查询条件", "与猎聘找人页面字段对应，例如关键词、城市、年限、学历、职位、行业。", "解析后逐项确认，没把握的条件先关闭"],
            ["评分阈值", "候选人达到该分数才进入自动沟通流程。", "初期建议偏高，跑出样本后再调"],
            ["默认话术", "可覆盖系统根据 JD 生成的补充话术。", "适合固定品牌口径或特殊岗位"],
        ],
        [1.35, 2.65, 2.35],
    )
    add_heading(doc, "JD 解析", 2)
    add_body(doc, "点击解析 JD 后，系统在后台调用千问，不会卡主界面。解析结果会进入条件确认弹窗，用户可选择哪些条件启用、哪些不用。")
    add_bullet(doc, "关键词适合放技能、业务场景和工具名，例如数字孪生、工业仿真、Unreal Engine、Omniverse、APS。")
    add_bullet(doc, "城市、职位、行业这类有限定选项的字段，以猎聘页面可选项为准；解析结果只是建议，不是强制值。")
    add_bullet(doc, "院校、专业、语言等条件如果过窄，会显著减少候选人池，建议只在必要时启用。")

    add_heading(doc, "5. 搜索条件写入与找人", 1)
    add_body(doc, "任务执行时会先进入猎聘找人页，再按岗位的查询条件逐项写入页面。城市和职位类条件会打开猎聘弹窗，系统按录制和页面结构选择选项并确认。")
    add_matrix(
        doc,
        ["条件类型", "系统处理方式", "验收点"],
        [
            ["关键词", "写入顶部搜索框，并选择关键词匹配方式。", "搜索框中能看到完整关键词"],
            ["工作年限 / 学历", "通过下拉或标签条件写入。", "页面下方标签出现，例如 5-10年、本科、硕士"],
            ["城市", "打开城市弹窗，选择城市并点击确认。", "标签区出现城市名称，而不是弹窗直接关闭"],
            ["职位 / 行业", "优先匹配猎聘有限选项；无法匹配时保守跳过。", "标签区出现职位或行业标签"],
            ["活跃度 / 性别等", "按页面已有选项选择。", "下拉框显示选中值"],
        ],
        [1.45, 2.75, 2.1],
    )
    add_note(doc, "频率控制", "不要随手把条件写入间隔改得很短，也不要连续高频点击搜索。保持系统默认节奏，降低账号风控风险。")

    add_heading(doc, "6. 任务管理", 1)
    add_body(doc, "任务是自动化的主入口。一个任务绑定一个账号和一个岗位，系统按照任务排序顺序逐个执行，不并发抢同一个右侧页面。")
    add_matrix(
        doc,
        ["任务字段", "作用", "建议"],
        [
            ["账号", "决定用哪个猎聘账号和 Profile 执行。", "按账号权益、岗位方向分配"],
            ["岗位", "决定 JD、查询条件、阈值和话术。", "一个岗位可创建多个不同账号任务"],
            ["候选人上限", "本任务最多处理多少人。", "测试 3-5，正式 30 起步再观察"],
            ["打招呼阈值", "达到阈值才执行沟通。", "可跟随岗位，也可任务单独覆盖"],
            ["AI 评分", "开启走千问；关闭走关键词规则评分。", "预算敏感或粗筛时关闭，精筛时开启"],
            ["自动沟通 / Dry-run", "决定是否真实发送。Dry-run 只生成或填入，不最终发送。", "新岗位先 Dry-run，稳定后再实发"],
            ["排序 / 优先级", "决定队列执行顺序。", "紧急岗位放前面"],
        ],
        [1.35, 2.4, 2.55],
    )
    add_heading(doc, "任务执行链路", 2)
    for step in [
        "切换到任务账号 Profile，检查登录状态。",
        "打开猎聘找人页，写入岗位查询条件并搜索。",
        "抓取当前页候选人列表摘要，保存到候选人清单。",
        "按当前页顺序打开候选人详情页，抓取完整简历文本。",
        "后台评分并生成补充话术；达到阈值且允许沟通时执行开聊。",
        "关闭当前候选人页面，打开下一位；当前页结束后点击下一页继续。",
        "达到候选人上限、无下一页、被终止或遇到人工处理节点时停止。",
    ]:
        add_number(doc, step)
    add_note(doc, "终止", "任务运行中使用“终止当前任务”。终止后后台步骤会在下一个检查点退出；如右侧页面已打开，系统不会再继续发送或翻页。")

    add_heading(doc, "7. 候选人清单", 1)
    add_body(doc, "候选人清单是最终资产，不只是调试结果。每条记录归属到岗位，并记录来源账号、来源任务、源网页、简历快照、评分、话术和沟通状态。")
    add_matrix(
        doc,
        ["功能", "说明"],
        [
            ["排序", "可按评分高到低、评分低到高、最近获取、最近更新排序。评分列在人名列后，便于快速扫人。"],
            ["筛选", "支持岗位筛选、评分区间筛选、多选状态筛选；可同时看未评分、已评分、已沟通、已淘汰等状态。"],
            ["淘汰 / 恢复", "不合适的人可以淘汰，后续用状态筛选隐藏；误淘汰可恢复。"],
            ["查看详情", "弹窗展示关键资料、评分报告、沟通记录，不占用右侧猎聘工作区。"],
            ["打开简历", "用抓取该候选人时的账号打开猎聘源网页，避免因为当前账号不一致跳登录页。"],
            ["查看全文快照", "展示当时抓到的简历正文和完整性信息，用于判断抓取质量。"],
        ],
        [1.45, 4.85],
    )
    add_note(doc, "数据口径", "列表卡片只提供摘要，最终候选人姓名、经历和评分应以简历详情页抓取结果为准。")

    add_heading(doc, "8. 评分与自动沟通", 1)
    add_body(doc, "评分分两种模式：AI 评分和关键词规则评分。AI 评分调用千问，适合精筛；关键词规则不走大模型，适合省 token 的粗筛。")
    add_matrix(
        doc,
        ["模式", "消耗", "适用场景"],
        [
            ["AI 评分", "会消耗大模型 token。", "岗位复杂、需要判断项目相关性和风险点。"],
            ["关键词规则", "不调用大模型。", "先扫候选池、预算敏感、只按必备词粗筛。"],
        ],
        [1.35, 1.7, 3.25],
    )
    add_heading(doc, "沟通流程", 2)
    add_body(doc, "达到阈值后，系统先点击“立即沟通”，在猎聘开聊弹窗中选择页面里的默认开聊语，然后进入聊天界面发送或填入系统生成的补充话术。")
    add_bullet(doc, "如果 Dry-run 开启：系统只做选择、生成或填入，不做最终发送。")
    add_bullet(doc, "如果 Dry-run 关闭：系统会尝试发送猎聘默认开聊语，并继续发送评分里生成的补充话术。")
    add_bullet(doc, "沟通状态会写入候选人和 greeting_logs，例如 generated、followup_filled_not_sent、followup_sent。")
    add_note(doc, "补充话术", "岗位里可维护自定义话术；没有自定义时，系统会根据 JD、候选人经历和评分结果生成。")

    add_heading(doc, "9. 调试、日志与录制", 1)
    add_body(doc, "默认界面保持清爽，调试信息收在调试按钮里。只有页面结构变化、抓取异常或需要单测功能时再打开。")
    add_matrix(
        doc,
        ["工具", "用途", "什么时候用"],
        [
            ["过程日志", "记录 URL、步骤、脚本结果、异常、候选人数量、沟通状态。", "任务卡住、统计异常、页面没跳转时"],
            ["抓当前简历", "直接从当前详情页抓取并保存快照。", "人工打开候选人后验证抓取质量"],
            ["查看抓取全文", "看最近一次抓到的正文。", "判断简历是否完整"],
            ["录制网页操作", "记录用户在猎聘页面上的真实操作路径。", "城市、职位、搜索按钮等页面结构变化时"],
            ["测试下一页", "只点击分页并验收页面变化，不跑 30 个候选人。", "验证翻页按钮是否能稳定命中"],
        ],
        [1.35, 2.3, 2.65],
    )
    add_heading(doc, "简历完整性判断", 2)
    add_bullet(doc, "日志里“是否详情页：是”说明当前页面被识别为简历详情页。")
    add_bullet(doc, "text_length、line_count 越高，通常抓取越完整；如果明显很短，要检查页面是否还停在列表页。")
    add_bullet(doc, "matched_sections 会显示求职意向、工作经历、项目经历、教育经历等命中情况。")
    add_bullet(doc, "附件简历会标记风险，系统不会自动索要附件，避免触发权益消耗或外部动作。")

    add_heading(doc, "10. 数据目录、迁移与 Windows 交付", 1)
    add_matrix(
        doc,
        ["路径", "内容", "处理方式"],
        [
            [".env", "千问/OpenAI Key 和模型配置。", "打包时会随成品复制；不要发给无关人员。"],
            ["data/app.db", "账号、岗位、任务、候选人、简历快照、评分、沟通和日志。", "迁移时复制；备份前确认合规。"],
            ["profiles/app", "每个猎聘账号的浏览器 Profile 和登录态。", "迁移可复用登录态，但非常敏感。"],
            ["data/recordings", "录制过的网页操作指纹。", "页面结构变化时可保留用于复现。"],
            ["dist/LiepinRecruitingAgent-win64.zip", "Windows 可分发压缩包。", "仅发给可信使用者。"],
        ],
        [1.6, 2.35, 2.35],
    )
    add_heading(doc, "Windows 打包", 2)
    add_body(doc, "在 Windows 10/11、Python 3.11 64 位环境下，进入项目根目录执行：")
    add_body(doc, "Set-ExecutionPolicy -Scope Process Bypass")
    add_body(doc, ".\\scripts\\build_windows.ps1 -Clean")
    add_body(doc, "完成后使用 dist\\LiepinRecruitingAgent-win64.zip；解压后双击 LiepinRecruitingAgent.exe。")
    add_note(doc, "API Key", "如果项目根目录有 .env，打包脚本会复制到成品目录。这样软件可直接使用千问能力，但包本身也因此包含敏感密钥。")

    add_heading(doc, "11. 常见问题", 1)
    add_matrix(
        doc,
        ["现象", "优先检查"],
        [
            ["打开后右侧跳登录页", "确认候选人或任务绑定的来源账号；必要时重新登录该账号。"],
            ["条件填完卡住", "看日志是否已点击搜索、是否页面弹窗未确认；用录制或单测定位具体组件。"],
            ["城市没生效", "确认城市弹窗里已选中并点击确认，标签区出现城市名称才算生效。"],
            ["任务停在搜索结果页", "看是否未打开候选人详情页；日志应出现 OPEN_RESUME 和“是否详情页：是”。"],
            ["只跑第一页", "使用“测试下一页”验证分页按钮；正常任务应当前页处理完再翻下一页。"],
            ["已沟通但统计没变", "看 greeting_status 是否为 followup_sent 或有效沟通状态，并刷新候选人列表。"],
            ["AI 不生效", "检查 .env 是否在运行目录，任务是否开启 AI 评分。"],
            ["界面卡顿", "评分、JD 解析应在后台执行；如持续卡顿，查看日志是否有异常堆积。"],
        ],
        [2.0, 4.3],
    )

    add_heading(doc, "12. 合规与安全提醒", 1)
    add_bullet(doc, "只在你有合法权限和业务需要的情况下处理候选人简历。")
    add_bullet(doc, "控制搜索、打开简历和沟通频率，避免对平台造成异常访问。")
    add_bullet(doc, "不要把包含 .env、data、profiles 的完整软件包发给无关人员。")
    add_bullet(doc, "真实发送前建议先跑 Dry-run，并人工抽查搜索条件、评分报告和话术。")
    add_bullet(doc, "附件简历、联系方式、验证码、账号安全提醒等节点应由人工确认。")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
